# -*- coding: utf-8 -*-
"""
Generira "05 - Prijedlog arhitekture.docx" - dokument za raspravu s direktorom
o nacinu izrade novog ERP/MES/WMS sustava (UI pristup, tehnologija, hosting).

Pokrenuti: python generate_arch_proposal.py
"""

import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from generate_docs import new_doc, add_table, bullets, numbered, OUT_DIR


def highlight_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
    return p


def doc_05():
    doc = new_doc(
        "Prijedlog arhitekture novog sustava",
        "Dokument za raspravu i odluku - opcije za izgled, tehnologiju i "
        "smjestaj (hosting) novog ERP/MES/WMS sustava. 15.06.2026.",
    )

    # -----------------------------------------------------------------
    doc.add_heading("1. Kontekst i cilj", level=1)
    doc.add_paragraph(
        "Cilj je izraditi novi ERP/MES/WMS sustav koji ce postupno zamijeniti "
        "trenutni ERP (legacy ERP). Plan je izgraditi cijeli sustav, a zatim prebaciti "
        "sve odjele (skladiste, proizvodnja, kvaliteta, nabava, prodaja...) "
        "odjednom."
    )
    doc.add_paragraph(
        "Prvi modul koji se gradi su zajednicki maticni podaci (artikli, "
        "kupci/dobavljaci, strojevi) - temelj na koji se nadograduju svi ostali "
        "moduli."
    )
    doc.add_paragraph(
        "Prije pocetka razvoja potrebno je odluciti tri stvari koje utjecu na "
        "SVE module: (1) kako ce aplikacija izgledati i pokretati se korisnicima, "
        "(2) koja tehnologija/programski jezik, (3) gdje ce sustav fizicki raditi "
        "(hosting). Ovaj dokument predlaze rjesenje za svaku stavku, s "
        "alternativama, za razgovor s direktorom."
    )

    # -----------------------------------------------------------------
    doc.add_heading("2. Pitanje 1: Kako ce aplikacija izgledati korisnicima?", level=1)
    doc.add_paragraph(
        "Direktor je izrazio zelju da aplikacija radi kao '.exe' program "
        "(dvoklik na ikonu, otvori se prozor programa), a ne kao web stranica "
        "(poput legacy ERP, koji se otvara u browseru). Postoje tri nacina da se to "
        "postigne, svaki s drugacijim omjerom brzine razvoja i krajnjeg dojma."
    )

    add_table(
        doc,
        ["Opcija", "Kako radi", "Prednosti", "Nedostaci"],
        [
            (
                "A) Web tehnologija + desktop prozor (pywebview)",
                "Aplikacija se razvija kao web app (kao 'Skladiste'), ali se "
                "pokrece u vlastitom prozoru bez adresne trake i menija "
                "preglednika - korisnik vidi samo program, ne 'browser'.",
                "Brz razvoj (gotove tablice, filteri, grafovi, izvjestaji); isti "
                "kod moze posluziti i na tabletu/mobitelu (npr. skeniranje u "
                "skladistu); nadovezuje se na postojece iskustvo (Skladiste 3.1).",
                "Tehnicki je 'web ispod haube' (korisnik to ne primjecuje); "
                "rijetke duboke Windows integracije (npr. system tray) traze "
                "dodatni rad.",
            ),
            (
                "B) Prava desktop aplikacija (npr. PyQt)",
                "Cijelo sucelje se piše izravno u desktop alatima (prozori, "
                "gumbi, tablice) bez ikakvog web servera.",
                "Najblize 'klasicnom' .exe osjecaju; najbolja integracija s "
                "Windowsom (ispis, dijaloski okviri, drag&drop).",
                "Znatno sporiji razvoj za ERP s desetcima ekrana (svaka tablica, "
                "filter, forma se rucno radi); ako se kasnije zatrazi pristup "
                "s tableta/mobitela, sucelje se mora raditi iznova.",
            ),
            (
                "C) Web aplikacija u browseru (kao legacy ERP)",
                "Aplikacija se otvara u Chrome/Edge/Firefox preko adrese na "
                "mrezi - isto kao legacy ERP danas.",
                "Najlakse za vise korisnika odjednom; najlakse za skladisne "
                "tablete i mobitele.",
                "Direktno ono sto direktor ne zeli - 'izgleda kao legacy ERP'.",
            ),
        ],
    )

    doc.add_heading("Preporuka za Pitanje 1", level=2)
    highlight_paragraph(doc, "Opcija A - web tehnologija u desktop prozoru (pywebview).")
    doc.add_paragraph(
        "Razlog: ERP/MES/WMS sa svim modulima (skladiste, proizvodnja, kvaliteta, "
        "nabava, prodaja...) je realno godine posla. Opcija B (prava desktop "
        "aplikacija) bi to produzila vise puta, jer se svaki ekran mora rucno "
        "graditi bez gotovih alata. Opcija A daje brzinu razvoja opcije C, ali "
        "krajnjem korisniku izgleda i pokrece se kao u opciji B (dvoklik na "
        "ikonu, prozor programa, bez vidljivog browsera). Dodatno, isti sustav "
        "kasnije moze posluziti skladisnom osoblju na tabletu (skeniranje QR "
        "kodova) bez ikakvog dodatnog razvoja."
    )

    # -----------------------------------------------------------------
    doc.add_heading("3. Pitanje 2: Koja tehnologija/programski jezik?", level=1)
    doc.add_paragraph(
        "Svi dosadasnji alati (Skladiste, Reklamacije, Normativi-skripte, "
        "predikcije potrosnje, generatori XML/PDF) napisani su u Pythonu. legacy ERP "
        "je napisan u .NET tehnologiji (ASP.NET)."
    )

    add_table(
        doc,
        ["Opcija", "Prednosti", "Nedostaci"],
        [
            (
                "Python (FastAPI + PostgreSQL) - preporuceno",
                "Logika iz postojecih alata (norme, predikcije, generiranje "
                "dokumenata, QR skladiste) se prilagodava i ponovno koristi, "
                "ne prepisuje od nule. Veliko postojece iskustvo.",
                "Nema direktne veze s legacy ERP (.NET) ekosustavom - ali to nije "
                "prepreka jer Python bez problema cita baze podataka (uklj. "
                "SQL Server kojeg legacy ERP vjerojatno koristi).",
            ),
            (
                ".NET / C#",
                "Isti ekosustav kao legacy ERP - moglo bi pomoci ako se kasnije "
                "duboko integrira s Paukovom bazom/kodom.",
                "Sav postojeci kod (13+ alata) bi se morao prepisati u novi "
                "jezik - veliki dodatni posao bez jasne dodane vrijednosti.",
            ),
        ],
    )

    doc.add_heading("Preporuka za Pitanje 2", level=2)
    highlight_paragraph(doc, "Python + FastAPI + PostgreSQL.")
    doc.add_paragraph(
        "Razlog: izbjegava prepisivanje vec gotove logike, a citanje Paukove "
        "baze (vjerojatno SQL Server) iz Pythona nije problem ako/kad se dobije "
        "pristup."
    )

    # -----------------------------------------------------------------
    doc.add_heading("4. Pitanje 3: Gdje sustav radi (hosting)?", level=1)
    doc.add_paragraph(
        "Postoji postojeci server na mrezi (npr. 'RSERVER'). Predlaze se da "
        "centralna baza podataka i 'mozak' sustava (backend) rade na tom "
        "serveru, a svako racunalo se spaja na njega preko mreze."
    )

    doc.add_heading("Skica rasporeda", level=2)
    add_table(
        doc,
        ["Lokacija", "Sto radi tamo"],
        [
            ("RSERVER (postojeci server)",
             "Centralna baza podataka (PostgreSQL) + 'mozak' sustava (backend) - "
             "jedan izvor istine za sve module i sve korisnike"),
            ("Uredska racunala",
             "'.exe' (desktop prozor) koji se spaja na RSERVER preko mreze - "
             "isti podaci, vise korisnika istovremeno"),
            ("Skladiste / proizvodnja (tableti, mobiteli)",
             "Browser na istu adresu - za skeniranje QR kodova, unos dnevnika "
             "rada na stroju i slicno, bez instalacije"),
        ],
    )

    doc.add_paragraph(
        "Ovakav raspored rjesava i 'puno korisnika odjednom' (svi se spajaju na "
        "isti RSERVER) i potrebu za skeniranjem/unosom na terenu (tableti u "
        "skladistu/proizvodnji), bez gradenja dva razlicita sustava."
    )

    # -----------------------------------------------------------------
    doc.add_heading("5. Sazetak prijedloga", level=1)
    bullets(
        doc,
        [
            "Korisnicki dozivljaj: '.exe' program na uredskim racunalima (Opcija A), "
            "+ browser na tabletima/mobitelima za skladiste i proizvodnju - isti "
            "sustav, bez dvostrukog razvoja.",
            "Tehnologija: Python + FastAPI + PostgreSQL - nastavak na postojeci kod "
            "i znanje, bez prepisivanja u novi jezik.",
            "Hosting: centralna baza i backend na RSERVER; svi klijenti (racunala, "
            "tableti) se spajaju preko mreze na isto mjesto.",
            "Prvi modul: zajednicki maticni podaci (artikli, kupci/dobavljaci, "
            "strojevi) - temelj za WMS, MES i Kvalitetu koji slijede.",
        ],
    )

    doc.add_heading("6. Otvoreno za odluku s direktorom", level=1)
    numbered(
        doc,
        [
            "Prihvaca li se pristup 'web u desktop prozoru' (Opcija A) kao nacin "
            "da aplikacija izgleda kao '.exe', ili direktor preferira pravu "
            "desktop aplikaciju (Opcija B) uz prihvacanje duzeg razvoja?",
            "Slaze li se Python + FastAPI + PostgreSQL, ili postoji razlog za "
            "razmatranje .NET-a (npr. postojeci IT ugovori, podrska)?",
            "Smije li se RSERVER koristiti za centralnu bazu i backend novog "
            "sustava, ili treba drugi server/racunalo?",
            "Treba li netko (IT/Vanado podrska) odobriti pristup Paukovoj bazi "
            "podataka za pocetni uvoz maticnih podataka (artikli, kupci, "
            "strojevi)?",
        ],
    )

    doc.save(os.path.join(OUT_DIR, "05 - Prijedlog arhitekture.docx"))


if __name__ == "__main__":
    doc_05()
    print("Gotovo:", os.path.join(OUT_DIR, "05 - Prijedlog arhitekture.docx"))
