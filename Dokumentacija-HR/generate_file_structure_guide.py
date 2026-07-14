# -*- coding: utf-8 -*-
"""
Generira hrvatski vodic (.docx) kroz strukturu datoteka i mapa ERP/MES/WMS projekta.
Pokrenuti: python generate_file_structure_guide.py
Izlaz: "07 - Struktura datoteka projekta.docx" u istoj mapi (Dokumentacija-HR).
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def new_doc(title, subtitle=None):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return doc


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph("")
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def mono(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)


def main():
    doc = new_doc(
        "Struktura datoteka projekta",
        "Vodic kroz mape i datoteke ERP/MES/WMS aplikacije - sto se gdje nalazi i "
        "cemu sluzi. Radna mapa: C:\\Users\\Tehnolog\\Desktop\\ERP-MES-WMS. "
        "Stanje: 15.06.2026.",
    )

    # ------------------------------------------------------------------
    doc.add_heading("1. Pregled na najvisoj razini", level=1)
    doc.add_paragraph(
        "Projekt je Python web-aplikacija (FastAPI). Sav izvorni kod nalazi se u "
        "mapi 'app', baza podataka i njene promjene u 'alembic', a dokumentacija u "
        "mapama 'docs' (tehnicka, engleski) i 'Dokumentacija-HR' (poslovna, "
        "hrvatski). Sljedeca tablica daje pregled svega na vrhu projektne mape."
    )
    add_table(
        doc,
        ["Mapa / datoteka", "Sto je to"],
        [
            ("app/", "Sav izvorni kod aplikacije - backend i HTML predlosci."),
            ("alembic/", "Migracije baze podataka - povijest svih promjena strukture baze, korak po korak."),
            ("docs/", "Tehnicka dokumentacija na engleskom (za razvoj) - shema baze, odluke, referenca na legacy ERP."),
            ("Dokumentacija-HR/", "Poslovna dokumentacija na hrvatskom (Word i PDF) - ovaj dokument je dio te mape."),
            ("Resources/", "Ulazni podaci za jednokratne uvoze, npr. Excel izvoz iz legacy ERP."),
            (".venv/", "Virtualno Python okruzenje - instalirane biblioteke. Generira se automatski, ne dira se rucno."),
            ("dev.db", "Lokalna baza podataka (SQLite) za razvoj/testiranje. Nije dio produkcije i nije u gitu."),
            ("alembic.ini", "Postavke alata za migracije baze (alembic)."),
            ("requirements.txt", "Popis Python biblioteka potrebnih za rad aplikacije."),
            ("README.md", "Kratak tehnicki uvod u projekt (engleski)."),
            ("legacy ERP acc.txt", "Pristupni podaci za legacy ERP ERP (samo za citanje). NIJE u gitu, ne smije se javno deliti."),
            (".gitignore", "Popis datoteka/mapa koje se ne spremaju u git (npr. dev.db, .venv, legacy ERP acc.txt)."),
            ("pokreni.bat", "Pokrece aplikaciju (dvoklik) i otvara je u pregledniku."),
            ("instalacija.bat", "Prva instalacija - postavlja Python okruzenje, biblioteke i bazu podataka."),
            ("azuriraj_bazu.bat", "Primjenjuje nove promjene baze podataka (migracije) na postojecu bazu."),
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("2. app/ - srce aplikacije", level=1)
    doc.add_paragraph(
        "Ovo je mapa s kodom koji se izvrsava. Sastoji se od ulazne tocke, "
        "zajednickih postavki, poslovnih modula i HTML predloska."
    )

    doc.add_heading("2.1 app/main.py - ulazna tocka", level=2)
    doc.add_paragraph(
        "Datoteka koja pokrece cijelu aplikaciju. Stvara FastAPI aplikaciju, "
        "registrira sve module (npr. Materijali, Pantoni) i postavlja da se "
        "pocetna stranica ('/') automatski preusmjeri na '/materijali'. Kad se "
        "doda novi modul (npr. Strojevi), ovdje se dodaje jedan red koji ga "
        "ukljucuje u aplikaciju."
    )

    doc.add_heading("2.2 app/core/ - zajednicke postavke", level=2)
    add_table(
        doc,
        ["Datoteka", "Sto radi"],
        [
            ("config.py", "Centralne postavke aplikacije - npr. na koju bazu podataka se spaja. "
                           "Lokalno koristi SQLite (dev.db), a na RSERVERU ce koristiti PostgreSQL "
                           "preko adrese definirane u .env datoteci."),
            ("database.py", "Uspostavlja konekciju na bazu podataka i pruza alat ('sesiju') koji "
                             "svaka stranica koristi za citanje/pisanje podataka."),
        ],
    )

    doc.add_heading("2.3 app/modules/ - poslovni moduli", level=2)
    doc.add_paragraph(
        "Svaki poslovni modul (Materijali, Pantoni, a u buducnosti Strojevi, "
        "Normativi...) ima vlastitu mapu unutar 'app/modules'. Svaka mapa modula "
        "obicno sadrzi do cetiri vrste datoteka:"
    )
    add_table(
        doc,
        ["Datoteka", "Uloga"],
        [
            ("models.py", "Definira strukturu podataka - koje tablice postoje u bazi i koja "
                           "polja (stupce) imaju. Ovo je 'izvor istine' za bazu podataka."),
            ("routes.py", "Definira web stranice i akcije modula - npr. popis, pregled/uredivanje "
                           "jednog zapisa, spremanje, brisanje. Ovdje je 'logika' modula."),
            ("schemas.py", "Opisuje oblik podataka za API/validaciju (koristi se djelomicno - "
                            "potpunije za buduce module s vanjskim sucelјem)."),
            ("import_*.py", "Jednokratne skripte za uvoz postojecih podataka (npr. iz legacy ERP Excel "
                             "izvoza) u novu bazu."),
        ],
    )

    doc.add_heading("Trenutni moduli", level=3)
    add_table(
        doc,
        ["Modul (mapa)", "Adresa u aplikaciji", "Sadrzaj"],
        [
            ("materijali/", "/materijali", "Sifrarnik materijala/artikala - popis, pregled, uredivanje, "
                                            "povijest cijena, uvoz iz legacy ERP (import_pauk.py)."),
            ("pantoni/", "/pantoni", "Sifrarnik Pantone boja (kod, naziv, hex boja za prikaz) - "
                                      "koristi se u materijalima i kasnije u Normativima."),
        ],
    )
    doc.add_paragraph(
        "Buduci moduli (Strojevi, Normativi i ostali iz plana razvoja) gradit ce se "
        "po istom obrascu - svaki dobiva vlastitu mapu s models.py / routes.py / "
        "schemas.py."
    )

    doc.add_heading("2.4 app/templates/ - izgled stranica (HTML)", level=2)
    doc.add_paragraph(
        "HTML predlosci koji odreduju kako stranice izgledaju. Koriste Jinja2 "
        "(predlosci s ugradenim Python izrazima), Tailwind CSS (stilovi), HTMX i "
        "Alpine.js (interaktivnost bez potpunog ponovnog ucitavanja stranice)."
    )
    add_table(
        doc,
        ["Datoteka / mapa", "Sto prikazuje"],
        [
            ("base.html", "Osnovni okvir svake stranice - zaglavlje, navigacijska traka (Materijali, "
                           "Pantoni...) i mjesto gdje se umece sadrzaj pojedine stranice."),
            ("materijali/list.html", "Popis svih materijala s tražilicom."),
            ("materijali/_table_body.html", "Samo tijelo tablice popisa - koristi se za brzo "
                                             "azuriranje rezultata pretrage (HTMX) bez ucitavanja "
                                             "cijele stranice."),
            ("materijali/_row.html", "Prikaz jednog reda tablice materijala."),
            ("materijali/detail.html", "Cijela kartica materijala - svi podaci grupirani u sekcije "
                                        "(osnovni podaci, klasifikacija, dimenzije, proizvodnja/boja "
                                        "s Pantone biracem, skladiste, cijene...)."),
            ("pantoni/list.html", "Popis Pantone boja s prikazom boje (swatch), dodavanje i brisanje."),
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("3. alembic/ - migracije baze podataka", level=1)
    doc.add_paragraph(
        "Svaka promjena strukture baze podataka (novo polje, nova tablica, nova "
        "veza) zapisuje se kao posebna 'migracija' - mala datoteka koja opisuje "
        "tocno sto se mijenja i kako se ta promjena moze poništiti. Migracije su "
        "lancano povezane (svaka znade koja je prethodna), tako da se baza moze "
        "korak po korak dovesti na najnoviju verziju naredbom iz 'azuriraj_bazu.bat'."
    )
    add_table(
        doc,
        ["Datoteka / mapa", "Sto je to"],
        [
            ("env.py, script.py.mako", "Konfiguracija alata za migracije (alembic) - ne mijenja se "
                                        "rucno u redovnom radu."),
            ("versions/", "Mapa sa svim migracijama, kronoloskim redom:"),
        ],
    )
    bullets(
        doc,
        [
            "63ce88be07d0 - stvaranje pocetnih tablica modula Materijali (materijal, "
            "materijal_papir, materijal_etiketa, cijena_povijest)",
            "4a6dd57b0dbf - dodavanje polja potrebnih za uvoz legacy ERP podataka",
            "833170c83337 - dodavanje preostalih legacy ERP polja u materijal",
            "90c84fcbe81a - dodatna legacy ERP polja (dimenzije, proizvodnja, boja...) i "
            "nova tablica 'pantone' za Pantone boje",
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("4. docs/ - tehnicka dokumentacija (engleski)", level=1)
    doc.add_paragraph(
        "Detaljna dokumentacija namijenjena razvoju - sadrzi nazive polja, sheme "
        "baze, tehnicke odluke. Pisana na engleskom."
    )
    add_table(
        doc,
        ["Datoteka", "Sadrzaj"],
        [
            ("01-domain-overview.md", "Pregled poslovanja - proizvodi, kupci, strojevi, tijek proizvodnje."),
            ("02-legacy-tools-catalog.md", "Katalog postojecih alata (Skladiste, Reklamacije, Normativi...)."),
            ("03-pauk-erp-reference.md", "Referenca trenutnog ERP-a (legacy ERP) - moduli, polja, obrasci."),
            ("04-open-decisions.md", "Arhitekturne i tehnoloske odluke."),
            ("05-modul-materijali.md", "Potpuna shema baze i UI za modul Materijali (uklj. Pantone i "
                                        "uvoz iz legacy ERP)."),
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("5. Dokumentacija-HR/ - poslovna dokumentacija (hrvatski)", level=1)
    doc.add_paragraph(
        "Dokumenti namijenjeni citanju bez tehnickog predznanja - u Word (.docx) i "
        "PDF formatu. Generiraju se Python skriptama (kako bi se lako mogli "
        "azurirati kad se nesto promijeni), a ne pisu se rucno u Wordu."
    )
    add_table(
        doc,
        ["Dokument", "Sadrzaj"],
        [
            ("00 - Pregled projekta", "Cilj projekta, trenutni status, struktura dokumentacije."),
            ("01 - Pregled poslovanja", "Djelatnost, materijali, strojevi, tijek proizvodnje."),
            ("02 - Katalog postojecih alata", "Svi postojeci alati i sto rade."),
            ("03 - Pregled legacy ERP ERP sustava", "Moduli i polja trenutnog ERP-a (legacy ERP)."),
            ("04 - Otvorena pitanja", "Pitanja/odluke koje treba razjasniti."),
            ("05 - Prijedlog arhitekture", "Prijedlog arhitekture novog sustava."),
            ("06 - Prijedlog tehnologije", "Prijedlog tehnologija (Python/FastAPI/...)."),
            ("07 - Struktura datoteka projekta", "Ovaj dokument - vodic kroz mape i datoteke."),
        ],
    )
    doc.add_paragraph("Skripte koje generiraju dokumente:")
    bullets(
        doc,
        [
            "generate_docs.py - generira dokumente 00-04",
            "generate_arch_proposal.py - generira dokument 05",
            "generate_tech_proposal.py - generira dokument 06",
            "generate_file_structure_guide.py - generira ovaj dokument (07)",
            "export_pdf.py - pretvara sve .docx dokumente u ovoj mapi u .pdf "
            "(koristeci instalirani Microsoft Word)",
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("6. Ostalo", level=1)
    add_table(
        doc,
        ["Datoteka", "Sto je to"],
        [
            ("Resources/Materijali(1).xlsx", "Izvoz sifrarnika materijala iz legacy ERP - koristi se kao "
                                              "ulazni podatak za jednokratni uvoz (import_pauk.py)."),
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("7. Kako se sve povezuje", level=1)
    numbered(
        doc,
        [
            "Pokretanje (pokreni.bat) podize aplikaciju definiranu u app/main.py.",
            "main.py ukljucuje module iz app/modules/ (npr. materijali, pantoni) - "
            "svaki modul dodaje svoje web adrese (npr. /materijali, /pantoni).",
            "Kad korisnik otvori stranicu, routes.py modula dohvati podatke iz baze "
            "(preko models.py i app/core/database.py) i prosljedi ih predlosku iz "
            "app/templates/, koji generira HTML prikaz.",
            "Kad se promijeni struktura podataka (models.py), generira se nova "
            "migracija u alembic/versions/ koja tu promjenu primjenjuje na bazu "
            "(azuriraj_bazu.bat).",
            "Dokumentacija u docs/ (tehnicka) i Dokumentacija-HR/ (poslovna) se "
            "rucno azurira kako se modul razvija.",
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("8. Pokretanje aplikacije - sazetak", level=1)
    bullets(
        doc,
        [
            "Prvi put: dvoklik na 'instalacija.bat' (postavlja okruzenje, biblioteke "
            "i bazu podataka).",
            "Svaki sljedeci put: dvoklik na 'pokreni.bat' (pokrece aplikaciju i "
            "otvara je u pregledniku na http://127.0.0.1:8000/materijali).",
            "Nakon preuzimanja novih promjena koda koje mijenjaju bazu: dvoklik na "
            "'azuriraj_bazu.bat'.",
        ],
    )

    doc.save(os.path.join(OUT_DIR, "07 - Struktura datoteka projekta.docx"))
    print("Spremljeno:", os.path.join(OUT_DIR, "07 - Struktura datoteka projekta.docx"))


if __name__ == "__main__":
    main()
