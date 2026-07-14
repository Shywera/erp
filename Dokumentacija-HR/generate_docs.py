# -*- coding: utf-8 -*-
"""
Generira hrvatsku, ljudima citljivu dokumentaciju (.docx) za ERP/MES/WMS projekt.
Pokrenuti: python generate_docs.py
Izlaz: .docx datoteke u istoj mapi (Dokumentacija-HR).
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def new_doc(title, subtitle=None):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.italic = True
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return doc


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph("")
    return table


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


# ---------------------------------------------------------------------------
# 00 - Pregled projekta
# ---------------------------------------------------------------------------
def doc_00():
    doc = new_doc(
        "ERP/MES/WMS Projekt - Pregled",
        "Novi jedinstveni informacijski sustav za tiskaru etiketa. Pokrenuto 15.06.2026.",
    )

    doc.add_heading("Cilj projekta", level=1)
    doc.add_paragraph(
        "Izraditi jedinstvenu ERP/MES/WMS aplikaciju za tiskaru etiketa, koja s vremenom "
        "objedinjuje i zamjenjuje desetke postojecih, odvojenih alata i Excel tablica "
        "razvijenih tijekom godina (mapa 'Arhiva skripta' na radnoj plohi)."
    )

    doc.add_heading("Trenutni status", level=1)
    doc.add_paragraph(
        "Faza istrazivanja i prikupljanja znanja. Jos nema napisanog koda nove "
        "aplikacije - prvo se gradi baza znanja o poslovanju, postojecim alatima i "
        "trenutnom ERP sustavu (legacy ERP), kako bi se donijele dobre arhitekturne odluke."
    )

    doc.add_heading("Struktura dokumentacije", level=1)
    doc.add_paragraph(
        "Radna mapa projekta: C:\\Users\\Tehnolog\\Desktop\\ERP-MES-WMS"
    )
    doc.add_paragraph("Dokumenti u ovoj (hrvatskoj) mapi:")
    bullets(
        doc,
        [
            "00 - Pregled projekta (ovaj dokument)",
            "01 - Pregled poslovanja - tvrtka, proizvodi, kupci, strojevi, tijek proizvodnje",
            "02 - Katalog postojecih alata - sto sve trenutno automatizirano postoji i sto radi",
            "03 - Pregled legacy ERP ERP sustava - sto trenutni ERP (legacy ERP) vec pokriva",
            "04 - Otvorena pitanja - odluke koje treba donijeti prije pocetka razvoja",
        ],
    )
    doc.add_paragraph(
        "Tehnicka (engleska, detaljnija) verzija dokumentacije nalazi se u mapi "
        "'docs' u istom projektu - namijenjena je razvoju, sadrzi putanje datoteka, "
        "nazive polja, formule i tehnicke detalje."
    )

    doc.add_heading("Napomena o pristupnim podacima", level=1)
    doc.add_paragraph(
        "Pristupni podaci za legacy ERP ERP (za citanje radi istrazivanja) nalaze se u "
        "datoteci 'legacy ERP acc.txt' u glavnoj projektnoj mapi. Ta datoteka nije dio "
        "verzioniranog koda (git) i ne smije se javno deliti."
    )

    doc.save(os.path.join(OUT_DIR, "00 - Pregled projekta.docx"))


# ---------------------------------------------------------------------------
# 01 - Pregled poslovanja
# ---------------------------------------------------------------------------
def doc_01():
    doc = new_doc(
        "Pregled poslovanja",
        "Djelatnost tvrtke, proizvodi, kupci, strojevi i tijek proizvodnje - kako je "
        "rekonstruirano iz postojecih alata i datoteka.",
    )

    doc.add_heading("Djelatnost", level=1)
    doc.add_paragraph(
        "Tvrtka proizvodi samoljepljive etikete na metaliziranom i bijelom papiru. "
        "Glavni kupci/klijenti:"
    )
    bullets(
        doc,
        [
            "kupac-pića - etikete s barkodovima (alat 'Cola' / BARGEN)",
            "kupac-moda - odjeca/logistika, EDI ASN XML dokumenti i etikete (alat 'Calz')",
            "Pivovare: Ozujsko pivo, 33 Export, Isenbeck",
            "ponavljajući ino-kupci - ponavljajuci upiti/ponude",
            "Switzerland WG Labels - cjenovni upiti",
            "Somersby, Paulaner - aktivni tenderi/ponude",
        ],
    )

    doc.add_heading("Materijali", level=1)
    bullets(
        doc,
        [
            "Metalizirani papir (Alu) i bijeli papir (Bijeli Pregani / Bijeli Glatki), "
            "razlicite gramature",
            "Plasticne folije - prate se po dimenzijama role (sirina mm x duljina m), "
            "pretvorba u kg preko konstante 0.01983 kg/m2",
            "Boje/tinte - cjenik se odrzava pod 'DOBAVLJAC-BOJA d.o.o.' (potrebno "
            "potvrditi je li to vlastita tvrtka ili dobavljac)",
        ],
    )

    doc.add_heading("Proizvodna opreme", level=1)
    doc.add_paragraph("Tisak (offset, Heidelberg):")
    bullets(
        doc,
        [
            "CX 104 6+LX (6 boja + lak)",
            "CX 102 5+LX (5 boja + lak)",
            "CD 102 6+LX (6 boja + UV lak)",
        ],
    )
    doc.add_paragraph("Doradni strojevi (rezanje, stancanje):")
    bullets(
        doc,
        [
            "Polar 137 (giljotina)",
            "SC20 (programabilna rezacica - rezanje traka, dorezivanje, pakiranje)",
            "MCS 115 (programabilna rezacica)",
            "Polar DC 11 (stancanje)",
            "Blumer Atlas 1110 Dual, Blumer Atlas 110 (stancanje)",
        ],
    )
    doc.add_paragraph(
        "Na pocetnoj stranici legacy ERP ERP-a prikazano je 6 'favorita' strojeva (ID: 23, "
        "37, 52, 57, 142, 143) - potrebno je utvrditi koji ID odgovara kojem stroju."
    )

    doc.add_heading("Tijek proizvodnje (kako je razumljeno do sada)", level=1)
    numbered(
        doc,
        [
            "Narudzba / tender / upit kupca (Excel kalkulacije, npr. 'Tender KUPAC.xlsx')",
            "Radni nalog u legacy ERP ('Narudzbe' / 'Radni nalozi')",
            "Planiranje proizvodnje (legacy ERP: 'Planiranje', 'Planiranje nabave iz proizvodnje')",
            "Tisak na offset stroju, zatim dorada/stancanje prema normativima "
            "(vidi alat 'Normativi' - formule za araka/h i etiketa/h po stroju)",
            "Evidencija rada po operaciji/stroju/djelatniku (legacy ERP: 'Dnevnici rada' / "
            "WorkOrderLog) - ovo su MES podaci: kolicine, senzorske kolicine, "
            "zavrsetak operacija",
            "Gotovi proizvodi na skladiste (palete s QR kodovima, raspored po "
            "lokacijama - prototip 'Skladiste')",
            "Otprema / dostava (legacy ERP: 'Otpremnice') - specificno po kupcu (npr. "
            "kupac-moda EDI ASN XML, kupac-pića etikete s barkodom)",
            "Fakturiranje (legacy ERP: 'Racuni')",
            "Kvaliteta - reklamacije i korektivne mjere (CAPA) - aplikacija 'Reklamacije'",
        ],
    )

    doc.add_heading("Otvorena pitanja o nazivima", level=1)
    bullets(
        doc,
        [
            "Je li 'DOBAVLJAC-BOJA d.o.o.' vlastita tvrtka ili dobavljac boja/tinte?",
            "Sto korisnicko ime u legacy ERP-u predstavlja u smislu odjela/uloge?",
        ],
    )

    doc.save(os.path.join(OUT_DIR, "01 - Pregled poslovanja.docx"))


# ---------------------------------------------------------------------------
# 02 - Katalog postojecih alata
# ---------------------------------------------------------------------------
def doc_02():
    doc = new_doc(
        "Katalog postojecih alata",
        "Pregled ~13 samostalnih alata/skripta razvijenih tijekom godina "
        "(C:\\Users\\Tehnolog\\Desktop\\Arhiva skripta). Ovo je glavni izvor znanja "
        "o tome sto novi sustav treba pokrivati.",
    )

    doc.add_heading("Pregledna tablica", level=1)
    add_table(
        doc,
        ["Alat", "Sto radi (ukratko)", "Buduci modul", "Zrelost"],
        [
            ("Skladiste 3.1", "QR skladiste, palete, FIFO/FEFO, web sucelje",
             "WMS - jezgra", "Najzrelije - vec web app (FastAPI)"),
            ("Reklamacije 1.0", "Reklamacije i korektivne mjere (CAPA), PDF/Excel/email",
             "Kvaliteta", "Najzrelije - puna Django aplikacija"),
            ("Tjedna usklada", "Tjedna provjera zaliha ispod minimuma, PDF izvjestaj",
             "WMS - narucivanje", "Skripta"),
            ("Utrosak vs stanje", "Usporedba dva stanja skladista, prepoznaje neaktivne zalihe",
             "WMS - izvjestaji", "Skripta + .exe"),
            ("Xgboost predikcije", "Predvidanje mjesecne potrosnje materijala (ML)",
             "MES - planiranje", "Skripta"),
            ("Normativi", "Norme rada strojeva (araka/h, etiketa/h)",
             "MES - norme", "Excel"),
            ("Montaza Etiketa", "Optimalan raspored etiketa na arku papira",
             "MES - planiranje", "Streamlit app"),
            ("Cola (BARGEN)", "Etikete s barkodom za kupac-pića narudzbe",
             "Otprema/logistika", "Skripta"),
            ("Calz (GENXML)", "EDI ASN XML i etikete za kupac-moda",
             "Otprema/logistika (EDI)", "Skripta"),
            ("Folije", "Izracun m2/kg folija iz naziva proizvoda",
             "Skladiste - izracuni", "Skripta/VBA"),
            ("Certifikati i teh.spec.", "Generiranje certifikata i tehnickih specifikacija",
             "Kvaliteta - dokumenti", "Prototip (mock API)"),
            ("Usklada metal", "Uskladivanje stanja metaliziranog papira (3 izvora)",
             "WMS/Nabava", "Skripta"),
            ("Update cjenika boja", "Masovno azuriranje cijena u cjeniku",
             "Nabava - cjenici", "Skripta + .exe"),
        ],
    )

    doc.add_heading("Skladiste 3.1 (WMS jezgra) - najzrelije", level=1)
    doc.add_paragraph(
        "Web aplikacija (FastAPI + SQLite) s prikazom rasporeda skladista (Plotly). "
        "Palete imaju QR kod s formatom 'QRID|DATUM_ULAZA|ROK|KOLICINA|SIFRA|LOT'. "
        "Skladiste ima oko 2000 pozicija na 18 regala (L, R1A-R8B, D, G, B), svaki s "
        "9-30 mjesta i 4-5 visina. Sustav automatski predlaze poziciju za novu paletu "
        "(prvo prazna mjesta u istom regalu na najnizoj visini, zatim u parnom regalu, "
        "zatim po prioritetu regala). Podrzava prijem (plan prijema -> dodjela pozicija "
        "-> skeniranje -> potvrda -> zatvaranje), izdavanje (skeniranje -> evidencija "
        "izlaza) i FIFO/FEFO pretrazivanje. Automatska sigurnosna kopija baze."
    )

    doc.add_heading("Reklamacije 1.0 (Kvaliteta/CAPA) - najzrelije", level=1)
    doc.add_paragraph(
        "Django web aplikacija za upravljanje reklamacijama (interne, kupca, "
        "dobavljaca) i korektivnim/preventivnim mjerama (CAPA)."
    )
    doc.add_paragraph("Reklamacija sadrzi: broj predmeta (RK-GGGG-NNNN), vrstu, status "
        "(NOVO -> U_OBRADI -> CEKA -> RIJESENO -> ZATVORENO), prioritet, kategoriju, "
        "naslov, opis, prijavitelja, kupca/dobavljaca, proizvod, broj radnog naloga, "
        "stroj, osoblje, datume, analizu uzroka (5 zasto), napomene i poveznicu na "
        "promjenu sustava (OB-21).")
    doc.add_paragraph("CAPA mjera sadrzi: vrstu (korektivna/preventivna), opis mjere, "
        "odgovornu osobu, rok, status, datum izvrsenja, rezultat i verifikaciju.")
    doc.add_paragraph("Izlazi: PDF izvjestaj, Excel izvoz (u bojama prema statusu), "
        "email obavijesti pri novoj reklamaciji i promjeni statusa.")

    doc.add_heading("Tjedna usklada", level=1)
    doc.add_paragraph(
        "Citanje izvjestaja 'IZVJESCE STANJA MATERIJALA', usporedba sa zadanim "
        "minimalnim kolicinama po artiklu (ukljucujuci posebna pravila za kombinirane "
        "artikle). Izlaz je PDF s crveno istaknutim artiklima koje treba naruciti."
    )

    doc.add_heading("Utrosak vs stanje", level=1)
    doc.add_paragraph(
        "Usporeduje dva snimka stanja skladista (npr. 05.01. vs 27.04.) i razvrstava "
        "artikle u kategorije: POVECANJE, TROSI SE, POTROSENO, NEAKTIVNI LAGER "
        "(nema kretanja, jos ima zaliha). Izlaz je Excel s bojama po statusu."
    )

    doc.add_heading("Xgboost predikcije potrosnje", level=1)
    doc.add_paragraph(
        "Model ucenja (XGBoost) predvida mjesecnu potrosnju materijala za 2026. na "
        "temelju povijesti 2023-2025. Koristi sezonske, trend i 'lag' znacajke. "
        "Izlaz: Excel s predikcijama po artiklu/mjesecu, validacija na stvarnim "
        "podacima i graf vaznosti znacajki."
    )

    doc.add_heading("Normativi", level=1)
    doc.add_paragraph(
        "Excel radna knjiga koja racuna norme rada (araka/h, etiketa/h) za strojeve "
        "tiska (CX104, CX102, CD102) i dorade (Polar 137, SC20, MCS115, Polar DC11, "
        "Blumer Atlas). Ulazni podaci: naklada, gramatura, format arka i etikete, "
        "broj boja, lakiranje, broj prolaza itd."
    )

    doc.add_heading("Montaza Etiketa", level=1)
    doc.add_paragraph(
        "Streamlit aplikacija koja za dvije vrste etiketa (A i B) pronalazi raspored "
        "stupaca na arku koji maksimizira broj kompletnih setova (A+B), uz minimalan "
        "otpad. Baza od 32 formata papira (3 vrste papira s razlicitim marginama)."
    )

    doc.add_heading("Cola (BARGEN) i Calz (GENXML)", level=1)
    doc.add_paragraph(
        "Cola: generira PDF etikete s Code128 barkodom (LOT, datum, kolicina) za "
        "kupac-pića narudzbe iz Excel tablice."
    )
    doc.add_paragraph(
        "Calz: generira EDI ASN XML dokumente (po kupac-moda shemi) i etikete za "
        "otpremu, na temelju Excel tablice s podacima o narudzbama, kolicinama, "
        "lotovima i barkodovima."
    )

    doc.add_heading("Folije", level=1)
    doc.add_paragraph(
        "Iz naziva proizvoda (npr. 'naziv 1000mm 500m') izracunava povrsinu (m2) i "
        "masu (kg) role folije, koristeci konstantu 0.01983 kg po m2. Postoji "
        "Python i VBA (Excel makro) verzija."
    )

    doc.add_heading("Certifikati i tehnicke specifikacije", level=1)
    doc.add_paragraph(
        "Generira Word/PDF certifikate i tehnicke specifikacije za radne naloge, "
        "popunjavajuci predloske s podacima o kupcu, materijalu i dimenzijama. "
        "Trenutno koristi simulirane (mock) podatke umjesto stvarnog ERP-a - ovo je "
        "tocka gdje bi se trebao spojiti na legacy ERP ili novi sustav."
    )

    doc.add_heading("Usklada metal", level=1)
    doc.add_paragraph(
        "Uskladuje stanje metaliziranog papira iz tri izvora: podloga nabave, "
        "izvjesce stanja materijala i raspored strojeva tiska. Izlaz je Excel s "
        "pivot tablicama po formatu, dobavljacu i mjesecu potrosnje."
    )

    doc.add_heading("Update cjenika boja", level=1)
    doc.add_paragraph(
        "Alat za masovno azuriranje cijena u cjeniku 'DOBAVLJAC-BOJA d.o.o.' (1630 "
        "stavki). Za svaki artikl pronalazi najnoviju cijenu, primjenjuje fiksno "
        "povecanje (EUR) ili postotak, i dodaje novi red s danasnjim datumom, "
        "cuvajuci povijest cijena."
    )

    doc.save(os.path.join(OUT_DIR, "02 - Katalog postojecih alata.docx"))


# ---------------------------------------------------------------------------
# 03 - Pregled legacy ERP ERP sustava
# ---------------------------------------------------------------------------
def doc_03():
    doc = new_doc(
        "Pregled legacy ERP ERP sustava",
        "Trenutni ERP sustav tvrtke, dobavljac Vanado (vanado.hr), pokrenut na "
        "lokalnom serveru https://erp.interno. Pregledano 15.06.2026, samo citanje "
        "(bez ikakvih izmjena).",
    )

    doc.add_heading("Opcenito", level=1)
    bullets(
        doc,
        [
            "Adresa: https://erp.interno (lokalna mreza)",
            "Tehnologija: ASP.NET MVC 5 (IIS, .NET 4.0), sucelje na hrvatskom jeziku",
            "Dobavljac: Vanado (vanado.hr)",
            "Pristupni podaci za citanje: 'legacy ERP acc.txt' u glavnoj projektnoj mapi "
            "(nije u git repozitoriju)",
        ],
    )

    doc.add_heading("Popis modula (izborni meni)", level=1)
    add_table(
        doc,
        ["Naziv u meniju", "Putanja", "Znacenje"],
        [
            ("Planiranje nabave iz proizvodnje", "/Plan/Supply", "Planiranje nabave na temelju proizvodnje"),
            ("Planiranje", "/Plan/Production", "Planiranje proizvodnje"),
            ("Materijal", "/Product/Index/1", "Sifrarnik materijala/artikala"),
            ("Upiti / Kalkulacije / sifrarnik", "/Product/Index2/4", "Upiti, kalkulacije, sifrarnik"),
            ("Knjiga izvoza", "/OutgoingBook", "Evidencija izvoza"),
            ("Knjiga uvoza", "/IncomingBook", "Evidencija uvoza"),
            ("Cjenici", "/Catalog/Sales", "Prodajni cjenici"),
            ("Cjenik nabava", "/Catalog", "Nabavni cjenik"),
            ("Ponude", "/Offer", "Ponude"),
            ("Narudzbe", "/Order", "Narudzbe"),
            ("Analiza prodaje", "/Warehouse/AnalysisBuyer", "Analiza prodaje po kupcu"),
            ("Radni nalozi", "/WorkOrder/Launched", "Radni nalozi (proizvodnja)"),
            ("Nabava", "/Request/Order", "Zahtjevi za nabavu"),
            ("Primke i medjuskladisnice", "/Purchase", "Primke robe i prijenosi izmedu skladista"),
            ("Ulazni racuni", "/IncomeBill/Index", "Ulazni racuni"),
            ("Otpremnice", "/Delivery", "Otpremnice"),
            ("Racuni", "/Order/Bill", "Izlazni (prodajni) racuni"),
            ("Proforma", "/Proforma/Bill", "Proforma racuni"),
            ("Skladiste", "/Warehouse/Index", "Stanje skladista"),
            ("Kategorije", "/Category", "Kategorije artikala"),
            ("Kupci / Dobavljaci", "/Contact", "Sifrarnik kupaca i dobavljaca"),
            ("Dnevnici rada", "/WorkOrderLog", "Evidencija rada - MES podaci"),
            ("Izvjestaji planova", "/PlanReports", "Izvjestaji o planovima"),
            ("Odrzavanje", "/Delay", "Odrzavanje strojeva"),
            ("Inventura", "/InventoryDocument", "Inventura skladista"),
        ],
    )

    doc.add_paragraph(
        "Na pocetnoj stranici prikazano je 6 'favorita' strojeva (ID 23, 37, 52, 57, "
        "142, 143) i popis trenutno otvorenih radnih naloga (vise od stotinjak, "
        "raspon ID-eva oko 64600-66300)."
    )

    doc.add_heading("Polja po modulu", level=1)

    doc.add_heading("Materijal (sifrarnik artikala)", level=2)
    doc.add_paragraph(
        "RB, Vrijeme izrade, Kategorija, Tip, Grupa, Podgrupa, Naziv, Sifra, "
        "Sipo sifra, Jedinica, Aktivno, Minimalna kolicina, Dobavljac, Datum zadnje "
        "izmjene, Zadnje izmjenio, Datoteke, excel predlozak."
    )
    doc.add_paragraph(
        "Ova polja (Kategorija/Tip/Grupa/Podgrupa) odgovaraju klasifikaciji koja se "
        "koristi u vecini postojecih alata (predikcije, usklade, itd.)."
    )

    doc.add_heading("Skladiste (stanje)", level=2)
    doc.add_paragraph(
        "RB, Sifra, Naziv proizvoda, Gramatura, Duljina [mm], Sirina [mm], Skladiste, "
        "Jedinica, raspolozivo po lotu, broj pakiranja, Minimalna kolicina skladista, "
        "Datum prvog ulaska, Datum zadnjeg ulaska, Datum zadnjeg izlaska, lot, "
        "prosjecna nabavna cijena, iznos, ukupna kolicina po sifri, ukupan iznos po "
        "sifri."
    )
    doc.add_paragraph(
        "Ovo je 'zivi' ekvivalent izvjestaja 'IZVJESCE STANJA MATERIJALA' koji koriste "
        "Tjedna usklada, Utrosak vs stanje, Usklada metal i Xgboost predikcije. "
        "Potvrdeno: legacy ERP je izvorni sustav za te izvjestaje."
    )

    doc.add_heading("Radni nalozi", level=2)
    doc.add_paragraph(
        "RB, Godina, Vrijeme izrade, Rok isporuke, Kupac, Oznaka RN, Sifra kupca, "
        "Broj RN, Unio, Napomena, Status, Datum zavrsetka, Kolicina [kom], Proizvodi, "
        "Oznaci."
    )

    doc.add_heading("Dnevnici rada (MES podaci)", level=2)
    doc.add_paragraph(
        "RB, Vrijeme izrade, Vrijeme zavrsetka, Broj radnog naloga, Kalkulacija, "
        "Proces, Artikl, Kolicina, Kolicina senzor, Djelatnik, Ime stroja, Operacija, "
        "Zavrsi operaciju, Zavrsi proizvod, Zavrsena faza na stroju, Zavrsena faza na "
        "stroju napomena, Zadnje izmjenio, Datum zadnje izmjene."
    )
    doc.add_paragraph(
        "Ovo je najblize postojecim MES podacima: evidencija po radnom nalogu, "
        "stroju, djelatniku i operaciji, s rucno unesenom i senzorski ocitanom "
        "kolicinom. Novi MES modul bi se trebao nadovezati na ovo ili ga zamijeniti "
        "boljim sustavom pracenja u stvarnom vremenu."
    )

    doc.add_heading("Kupci / Dobavljaci", level=2)
    doc.add_paragraph(
        "RB, Sifra, Naziv, Interni naziv, OIB, Mjesto, Maticni broj, Valuta placanja "
        "[dan], Referent, Radno vrijeme, Grupa, Tip kupca, HBOR osiguranje, HBOR "
        "ugovoreni rok placanja [dan], Naziv, excel predlozak."
    )

    doc.add_heading("Sto to znaci za novi sustav", level=1)
    bullets(
        doc,
        [
            "legacy ERP vec pokriva: sifrarnik materijala, kupce/dobavljace, radne naloge, "
            "osnovnu evidenciju rada, stanje skladista, cjenike, fakturiranje, "
            "otpremnice, knjige uvoza/izvoza, odrzavanje i inventuru.",
            "Postojeci alati postoje upravo zato sto legacy ERP ne radi: napomene o "
            "narucivanju, predikcije potrosnje, uskladivanje, QR skladiste, "
            "optimizaciju rasporeda, EDI/specificne etikete za kupce i kvalitetu/CAPA.",
            "Dvije moguce arhitekture (jos nije odluceno): (1) 'satelitski' sustav koji "
            "cita/pise u legacy ERP (preko izvoza ili API-ja ako postoji) i pokriva praznine, "
            "ili (2) novi sustav postaje glavni izvor istine, s postupnom migracijom "
            "modula iz legacy ERP.",
            "Treba provjeriti ima li legacy ERP JSON/API sloj iznad HTML stranica - to bi "
            "znacajno olaksalo integraciju.",
        ],
    )

    doc.save(os.path.join(OUT_DIR, "03 - Pregled legacy ERP ERP sustava.docx"))


# ---------------------------------------------------------------------------
# 04 - Otvorena pitanja
# ---------------------------------------------------------------------------
def doc_04():
    doc = new_doc(
        "Otvorena pitanja / odluke",
        "Pitanja koja treba razjasniti prije konacnih arhitekturnih odluka. "
        "Azurirati ovaj dokument kako pristizu odgovori.",
    )

    doc.add_heading("Na cekanju", level=1)
    numbered(
        doc,
        [
            "Odnos prema legacy ERP - zamjenjuje li novi sustav legacy ERP, radi paralelno "
            "(uvoz izvjestaja), ili se odlucuje po modulu?",
            "Ima li legacy ERP API osim HTML stranica? Ako da, integracija je puno "
            "jednostavnija.",
            "Tehnologija za novi sustav - kandidati: FastAPI + SQLite + web "
            "(kao Skladiste 3.1) ili Django + SQLite/Postgres (kao Reklamacije 1.0). "
            "Korisnik je trazio preporuku kada se prikupi vise informacija.",
            "Koji modul prvi - WMS (nadogradnja Skladiste 3.1), zajednicki "
            "matricni podaci (artikli/kupci/strojevi), Kvaliteta/CAPA "
            "(Reklamacije 1.0), ili Proizvodnja/MES (normativi + planiranje)?",
            "Model koristenja - jedan korisnik ili vise korisnika/cijela tvrtka "
            "(Skladiste 3.1 vec radi kao mrezni server, sto upucuje na vise "
            "korisnika)?",
            "Je li 'DOBAVLJAC-BOJA d.o.o.' vlastita tvrtka ili dobavljac?",
            "Mapiranje ID-eva strojeva u legacy ERP (23, 37, 52, 57, 142, 143) na "
            "stvarne nazive strojeva (CX104, CX102, CD102, Polar, SC20, MCS115, "
            "Blumer Atlas...).",
        ],
    )

    doc.add_heading("Odluceno", level=1)
    doc.add_paragraph("(jos nista)")

    doc.save(os.path.join(OUT_DIR, "04 - Otvorena pitanja.docx"))


if __name__ == "__main__":
    doc_00()
    doc_01()
    doc_02()
    doc_03()
    doc_04()
    print("Gotovo. Datoteke spremljene u:", OUT_DIR)
