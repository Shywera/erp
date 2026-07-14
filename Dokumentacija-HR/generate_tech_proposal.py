# -*- coding: utf-8 -*-
"""
Generira "06 - Prijedlog tehnologije.docx" - objektivna usporedba Python
(FastAPI+PostgreSQL+HTMX) stacka vs .NET, plus usporedba s trenutnim legacy ERP ERP-om.
Namjerno NE uzima u obzir postojeci kod/skripte - cisto tehnicka analiza.

Pokrenuti: python generate_tech_proposal.py
"""

import os
from docx.shared import RGBColor

from generate_docs import new_doc, add_table, bullets, numbered, OUT_DIR


def highlight_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
    return p


def doc_06():
    doc = new_doc(
        "Prijedlog tehnologije: Python vs .NET",
        "Cisto tehnicka usporedba (bez obzira na postojeci kod/skripte) + "
        "usporedba s trenutnim legacy ERP ERP sustavom. 15.06.2026.",
    )

    # -----------------------------------------------------------------
    doc.add_heading("Sazetak", level=1)
    highlight_paragraph(
        doc,
        "Preporuka: Python (FastAPI + PostgreSQL + HTMX/Alpine.js/Tailwind).",
    )
    doc.add_paragraph(
        "Ova preporuka je dana na temelju cisto tehnickih kriterija (brzina "
        "razvoja, troskovi, performanse, buduce potrebe), bez obzira na to sto "
        "vec postoji od koda. .NET je realna i kvalitetna alternativa - "
        "razlike su manje nego sto se cini, ali za ovaj tip projekta (interni "
        "ERP/MES/WMS s buducim modulima planiranja/predikcija) Python ima "
        "vise prednosti koje se gomilaju kroz vrijeme."
    )

    # -----------------------------------------------------------------
    doc.add_heading("1. Prijedlog stacka (detaljno)", level=1)
    add_table(
        doc,
        ["Dio sustava", "Tehnologija", "Zasto"],
        [
            ("Baza podataka", "PostgreSQL",
             "Besplatna, robusna, podrzava vise korisnika istovremeno, "
             "industrijski standard za poslovne aplikacije."),
            ("Backend (API i logika)", "FastAPI (Python)",
             "Moderan, asinkroni web framework. Automatski generira API "
             "dokumentaciju. Provjera tipova podataka (Pydantic) smanjuje "
             "greske."),
            ("Veza s bazom (ORM)", "SQLAlchemy + Alembic",
             "Standardni nacin rada s bazom iz Pythona; Alembic vodi povijest "
             "promjena strukture baze (migracije) - sigurno nadogradivanje."),
            ("Frontend (sucelje)", "Jinja2 + HTMX + Alpine.js",
             "Stranice se generiraju na serveru (brzo, jednostavno), a HTMX "
             "osvjezava samo dijelove stranice - osjeca se brzo kao moderna "
             "aplikacija, bez tezine React/Vue alata."),
            ("Izgled", "Tailwind CSS",
             "Moderan, cist izgled bez puno rucnog CSS-a."),
            ("Hosting", "RSERVER (postojeci server)",
             "Backend + baza rade na jednom mjestu, svi korisnici (racunala, "
             "tableti) pristupaju preko mreze."),
        ],
    )

    # -----------------------------------------------------------------
    doc.add_heading("2. Python (FastAPI) vs .NET (ASP.NET Core) - usporedba", level=1)
    doc.add_paragraph(
        "Usporedba pretpostavlja MODERNE verzije oba ekosustava: FastAPI "
        "(Python 3.12) i ASP.NET Core 8/9 (C#) - ne stariju ASP.NET MVC 5 "
        "tehnologiju koju koristi legacy ERP."
    )

    add_table(
        doc,
        ["Kriterij", "Python (FastAPI)", ".NET (ASP.NET Core)", "Komentar"],
        [
            (
                "Brzina razvoja CRUD ekrana (artikli, kupci, radni nalozi...)",
                "Vrlo brzo - kratak kod, manje 'ceremonije', puno gotovih "
                "biblioteka.",
                "Brzo, ali vise pocetnog koda po ekranu (klase, DTO-ovi, "
                "konfiguracija).",
                "Prednost: Python",
            ),
            (
                "Performanse (brzina obrade zahtjeva)",
                "Vrlo dobre za web API (async); dovoljne za desetke/stotine "
                "korisnika u tvrtki.",
                "Generalno najbrzi (kompajlirani jezik) - prednost kod "
                "ekstremnog opterecenja.",
                "Prednost: .NET, ali razlika se prakticki ne osjeca za "
                "interni ERP ovog obima.",
            ),
            (
                "Sigurnost tipova (type safety)",
                "Dobra uz Pydantic + type hints, ali nije ugradeno u jezik.",
                "Izvrsna - C# je strogo tipiziran jezik, greske se hvataju "
                "prije pokretanja.",
                "Prednost: .NET (manja razlika u praksi uz disciplinu)",
            ),
            (
                "Troskovi licenci",
                "0 EUR - sve (Python, FastAPI, PostgreSQL) je besplatno i "
                "open-source.",
                "ASP.NET Core je besplatan; ali ako se koristi SQL Server "
                "(umjesto PostgreSQL), licence mogu kostati. Visual Studio "
                "Community je besplatan.",
                "Prednost: Python (ako se i .NET kombinira s PostgreSQL-om, "
                "razlika je manja)",
            ),
            (
                "Multiplatformnost (Windows/Linux server)",
                "Radi identicno na Windows i Linux serveru.",
                "ASP.NET Core takoder radi na Linuxu, ali ekosustav je "
                "tradicionalno Windows-centricniji.",
                "Slicno - oba moderna, blaga prednost Python za jednostavnost "
                "na bilo kojem serveru.",
            ),
            (
                "Buduci moduli: predikcije, optimizacija, planiranje (MES)",
                "Vrlo jak ekosustav (pandas, scikit-learn, XGBoost, OR-Tools "
                "za optimizaciju rasporeda).",
                "ML.NET postoji, ali je znatno manje razvijen i manje "
                "koristen od Python alata.",
                "Prednost: Python - relevantno za module predikcije "
                "potrosnje, optimizacije rasporeda strojeva itd.",
            ),
            (
                "Generiranje dokumenata (PDF, Excel, Word, XML)",
                "Odlicne biblioteke (reportlab, openpyxl, python-docx, lxml).",
                "Takoder dobre biblioteke (iTextSharp/QuestPDF, "
                "ClosedXML, OpenXML SDK).",
                "Podjednako - oba pokrivena.",
            ),
            (
                "Dostupnost developera / znanje na trzistu",
                "Python je medu najpopularnijim jezicima opcenito, lako se "
                "uci.",
                "C#/.NET je vrlo zastupljen u poslovnim aplikacijama, "
                "pogotovo u Hrvatskoj (mnogi ERP-ovi su .NET, kao legacy ERP).",
                "Podjednako - ovisi o trzistu; .NET developera moze biti "
                "lakse naci lokalno za buduce poslovne aplikacije.",
            ),
            (
                "Alati za razvoj (IDE, debug, testiranje)",
                "Vrlo dobri (VS Code, PyCharm), ali manje 'integrirano' nego "
                "Visual Studio.",
                "Visual Studio je vrhunski integriran IDE - najbolji "
                "debugging i refaktoriranje.",
                "Blaga prednost: .NET",
            ),
            (
                "Odrzavanje na duzi rok (5-10 godina)",
                "Stabilno, ogromna globalna zajednica, jezik se ne mijenja "
                "drasticno.",
                "Stabilno, Microsoft podrska, redovita nadogradnja.",
                "Podjednako",
            ),
        ],
    )

    # -----------------------------------------------------------------
    doc.add_heading("3. Usporedba s trenutnim legacy ERP ERP-om", level=1)
    doc.add_paragraph(
        "legacy ERP koristi ASP.NET MVC 5 na .NET Framework 4.0 (tehnologija iz "
        "~2012. godine), server-rendered stranice s jQuery, pokrece se na IIS "
        "(samo Windows). Bez obzira koju tehnologiju odaberemo za novi sustav, "
        "ona ce biti znacajno modernija od legacy ERP."
    )

    add_table(
        doc,
        ["Aspekt", "legacy ERP (trenutni, ASP.NET MVC 5 / .NET 4.0)",
         "Novi sustav - Python prijedlog", "Novi sustav - .NET prijedlog (moderni)"],
        [
            ("Godina tehnologije", "~2012 (.NET Framework 4.0, MVC 5)",
             "2024+ (FastAPI, Python 3.12)", "2024+ (ASP.NET Core 8/9)"),
            ("Frontend", "Server-rendered + jQuery, pun reload stranica",
             "HTMX - djelomicna azuriranja, osjeca se brze",
             "Blazor ili Razor + HTMX - takoder moguce djelomicna azuriranja"),
            ("API/integracije", "Nije poznato da postoji API (samo HTML)",
             "FastAPI automatski generira API (OpenAPI/Swagger) - lako za "
             "buduce integracije (npr. skener, mobilna app)",
             "ASP.NET Core takoder lako generira API"),
            ("Platforma", "Windows/IIS samo",
             "Windows ili Linux server", "Windows ili Linux server (Core je "
             "multiplatformski)"),
            ("Licence", "Vanado licenca (nepoznat trosak)",
             "0 EUR (sve open-source)", "0 EUR za framework; baza ovisi o "
             "odabiru (PostgreSQL=0, SQL Server=moguci trosak)"),
        ],
    )

    doc.add_paragraph(
        "Zakljucak ove tocke: i Python i moderni .NET bi predstavljali veliki "
        "korak unaprijed u odnosu na legacy ERP. Odabir izmedu njih ne ovisi o "
        "legacy ERP, vec o kriterijima iz tocke 2."
    )

    # -----------------------------------------------------------------
    doc.add_heading("4. Zakljucak i preporuka", level=1)
    highlight_paragraph(doc, "Python: FastAPI + PostgreSQL + HTMX/Alpine.js/Tailwind")
    doc.add_paragraph("Razlozi (cisto tehnicki, neovisno o postojecem kodu):")
    bullets(
        doc,
        [
            "Najbrzi razvoj CRUD/izvjestajnih ekrana, kojih ce u ERP/MES/WMS "
            "biti puno (artikli, kupci, radni nalozi, skladiste, kvaliteta...).",
            "Nula troskova licenci - cijeli stack je besplatan i open-source.",
            "Najjaci ekosustav za module koji ce sigurno biti potrebni "
            "kasnije: predikcija potrosnje, optimizacija rasporeda "
            "proizvodnje, planiranje.",
            "HTMX pristup daje moderan, brz osjecaj korisniku bez tezine "
            "punog JavaScript frontend frameworka.",
            "Razlika u sirovim performansama prema .NET-u je u praksi "
            "nevidljiva za broj korisnika koji ce ovaj sustav imati "
            "(desetci, ne milijuni zahtjeva).",
        ],
    )
    doc.add_paragraph(
        ".NET ostaje legitimna alternativa ako se zeli najveca sigurnost "
        "tipova, najbolji IDE (Visual Studio) ili ako tvrtka iz drugih razloga "
        "(npr. postojece IT ugovore, lokalnu dostupnost .NET developera) "
        "zeli ostati u tom ekosustavu - ali za ovaj projekt, sam izbor "
        "tehnologije ne donosi prednost koja bi nadmasila Python prednosti "
        "iznad."
    )

    doc.save(os.path.join(OUT_DIR, "06 - Prijedlog tehnologije.docx"))


if __name__ == "__main__":
    doc_06()
    print("Gotovo:", os.path.join(OUT_DIR, "06 - Prijedlog tehnologije.docx"))
